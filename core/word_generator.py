import os
from itertools import zip_longest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def _set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def _set_cell_shading(cell, fill_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def create_translated_word_doc(pages_data: list[tuple[str, str]], output_path: str = "translated.docx"):
    doc = Document()
    
    # Set Narrow Margins (0.5 inches all around) and Center page vertically
    for section in doc.sections:
        section.top_margin = Cm(0.7)
        section.bottom_margin = Cm(0.7)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        
        # Vertically center page content
        sectPr = section._sectPr
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        sectPr.append(vAlign)
    
    total_pages = len(pages_data)
    
    for page_idx, (left_text, right_text) in enumerate(pages_data):
        # We create a table for this specific page
        table = doc.add_table(rows=0, cols=2)
        # 'Normal Table' has no borders by default in Word
        table.style = 'Normal Table'
        
        # Split texts by newline to match them line-by-line
        left_lines = left_text.split('\n') if left_text else []
        right_lines = right_text.split('\n') if right_text else []
        
        for row_idx, (l_line, r_line) in enumerate(zip_longest(left_lines, right_lines, fillvalue="")):
            row = table.add_row()
            row.height = Cm(0.52)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            row_cells = row.cells
            
            # --- Left Column (Korean, Black, Right-aligned) ---
            cell_left = row_cells[0]
            _set_cell_border(cell_left)
            p_left = cell_left.paragraphs[0]
            p_left.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            if l_line:
                run_left = p_left.add_run(l_line.strip())
                run_left.font.name = "Noto Sans KR Regular"
                run_left._element.rPr.rFonts.set(qn('w:eastAsia'), "Noto Sans KR Regular")
                run_left.font.color.rgb = RGBColor(0, 0, 0) # Black
                run_left.font.size = Pt(9)            
            # --- Right Column (English, Blue, Left-aligned) ---
            cell_right = row_cells[1]
            _set_cell_border(cell_right)
            
            # Alternate row background color for the English side to improve readability
            if row_idx % 2 != 0:
                _set_cell_shading(cell_right, "F2F2F2") # Light Gray
            
            p_right = cell_right.paragraphs[0]
            p_right.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            if r_line:
                run_right = p_right.add_run(r_line.strip())
                run_right.font.name = "Noto Sans KR Regular"
                run_right._element.rPr.rFonts.set(qn('w:eastAsia'), "Noto Sans KR Regular")
                run_right.font.color.rgb = RGBColor(0, 112, 192) # Blue
                run_right.font.size = Pt(9)
                
        # Add a page break after the table, unless it's the very last page
        if page_idx < total_pages - 1:
            doc.add_page_break()
            
    doc.save(output_path)
    return output_path
